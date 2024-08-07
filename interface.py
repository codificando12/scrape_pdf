from pypdf import PdfReader
from read_folder import read_dir
from docx import Document
from tkinter import filedialog
import tkinter as tk
from tkinter import *
import unicodedata
import re
import os

counter = 0

folder_path = "" #save the folder path that search folder path returns

save_folder_path = ""

def interface():
    
    def get_folder_path(): #This funtion return the folder path to get the pdfs list
        global folder_path
        folder_path = filedialog.askdirectory() 
        path_label.configure(text = folder_path)

    def save_file_path():
        global save_folder_path
        save_folder_path = filedialog.askdirectory()
        save_path_label.configure(text = save_folder_path) #this will save the folder path in a label

    # this function will get the keyword and proccess to be easy to look for it
    def get_text():
        search_word = search_box.get()
        word_normalized = unicodedata.normalize("NFKD", search_word).encode("ascii", "ignore").decode("ascii")
        return word_normalized

    #this one will get the file name
    def file_name():
        choose_file_name = file_name_box.get()
        return choose_file_name

    #this one will start the program
    def start_program():
        global folder_path
        global file_list
        global counter
        text_area.delete("1.0", END)
        file_list = read_dir(folder_path)
        print(len(file_list))
        lenth = len(file_list)
        counter = lenth
        file_counter_label.config(text = counter)
        window.update()
        print(lenth)
        book_page_paragraph = []
        word = get_text()
        compound_word = word.split(" ")
        # print(compound_word)
        file_name_choosen = file_name()
        word_document = Document()
        word_document.add_paragraph(f'Word = {" ".join(compound_word)}\n')

        submit_button.configure(state = tk.DISABLED)
        browse_button.configure(state = tk.DISABLED)
        save_browse_button.configure(state = tk.DISABLED)

        for file in range(len(file_list)):

            try:
                reader = PdfReader(f'{folder_path}/{file_list[file]}') 
                number_of_pages = len(reader.pages)
                print(file_list[file])
                text_area.insert(INSERT, f'{file_list[file]},\nTotal pages: {number_of_pages} \n ---------\n')
                text_area.see(END)
                print(number_of_pages)
            except:
                print(f" {file_list[file]}. La propiedad no está presente en el objeto o tiene un valor nulo.")
                error_text_area.insert(INSERT, f'The book "{file_list[file]}" couldn\'t be opened.\n ---------\n')
                error_text_area.see(END)
                # text_area.uptdate()
                continue
            
            for page_number in range(number_of_pages):

                try:    
                    page = reader.pages[page_number]
                    try:
                        text = page.extract_text()
                        text = text.lower()
                        text_normalized = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
                        text_normalized = text_normalized.replace("(", "( ")
                        text_normalized = text_normalized.replace("'", "' ")
                        text_normalized = text_normalized.replace("\"", "\" ")
                        text_normalized = text_normalized.replace("[", "[ ")
                        text_normalized = text_normalized.replace("{", "{ ")
                        text_normalized = text_normalized.replace(",", " ,")
                        text_normalized = text_normalized.replace(".", " .")
                        text_normalized = text_normalized.replace(";", " ;")
                        text_normalized = text_normalized.replace(":", " :")
                        paragraphs = text_normalized.split(' ')
                        # print(paragraphs)
                        
                    except:
                        continue
                except:
                    print(f"No se pudo leer las paginas en el libro '{file_list[file]}'")
                    pages_error_text_area.insert(INSERT, f'The pages of "{file_list[file]}" couldn\'t be readed.\n --------- \n')
                    pages_error_text_area.see(END)
                    continue
                
                for i in range(len(paragraphs)):
                    details = []
                    count = 0
                    if len(compound_word) == 1 and paragraphs[i].startswith(word):
                        details.append(f'---{file_list[file].upper()} ---')
                        details.append(f'---p.{page_number + 1}/ ---')
                        sentences = paragraphs[i - 50:i + 50]
                        details.append(f'---\n/{" ".join(sentences)}/ ---')
                        # print(sentences)
                        # print(details)
                        book_page_paragraph.append(details)
                        details = []
                        # print(details)
                    elif i + len(compound_word) <= len(paragraphs) and all(paragraphs[i + j].startswith(compound_word[j]) for j in range(len(compound_word))): 
                        
                        details.append(f'---{file_list[file].upper()} ---')
                        details.append(f'---p.{page_number + 1}/ ---')
                        sentences = paragraphs[i - 50:i + 50]
                        details.append(f'---\n/{" ".join(sentences)}/ ---')
                        book_page_paragraph.append(details)
                        details = []
                        count += 1          
                    else:
                        continue
        # word_document = Document()
            counter -= 1
            file_counter_label.config(text = counter)
            file_counter_label.update()
            for item in range(len(book_page_paragraph)):
                data = book_page_paragraph[item]
                joined_data = "".join(data)
                normalize_data = unicodedata.normalize("NFKD", joined_data).encode("ascii", "ignore").decode("ascii")
                clean_data = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', '', normalize_data)
                word_document.add_paragraph(clean_data)

            save_path = os.path.join(save_folder_path, f"{file_name_choosen}.docx")
            word_document.save(save_path)
            book_page_paragraph = []

        print("Scan completed")
        counter = 0
        text_area.insert(INSERT, "Scan Completed")
        file_counter_label.config(text = counter)
        window.update()
        submit_button.configure(state = tk.NORMAL)
        browse_button.configure(state = tk.NORMAL)
        save_browse_button.configure(state = tk.NORMAL)

    file_list = []

    window = tk.Tk()
    window.title("Ciappi PDF Scraper")
    window.geometry("1050x800")

    title_label = tk.Label(window, text="Search the word that you want in your PDFs files \n and get a paragraph with the results")
    title_label.grid(row=0, column=1)

    folder_label = tk.Label(window, text="Select PDF \nFolder")
    folder_label.grid(row=1, column=0)

    path_label = tk.Label(window, bg = "White", borderwidth = 1, relief = "solid" ,width = 50)
    path_label.grid(row = 1, column = 1,pady = 6)

    browse_button = tk.Button(window, text= "Browse", command = get_folder_path)
    browse_button.grid(row = 1, column = 2, padx = 6)

    search_label = tk.Label(window, text="Insert Word to Search \n(No Brackets)")
    search_label.grid(row = 2, column = 0)

    search_box = tk.Entry(window, width=58)
    search_box.grid(row = 2, column = 1, pady = 10)
    
    file_name_label = tk.Label(window, text="Choose File Result Name")
    file_name_label.grid(row = 3, column=0)

    file_name_box = tk.Entry(window, width = 58)
    file_name_box.grid(row = 3, column = 1, pady = 10)

    save_folder_label = tk.Label(window, text = "Save in Folder")
    save_folder_label.grid(row = 4, column = 0)

    save_path_label = Label(window, bg="White", width=50)
    save_path_label.grid(row = 4, column = 1)

    save_browse_button = Button(window, text="Browse", command = save_file_path)
    save_browse_button.grid(row = 4, column = 2, padx = 6)

  
    submit_button = tk.Button(window, text="Submit", width = 15, command = start_program)
    submit_button.grid(row = 5, column = 1)
    
    progres_label = tk.Label(window, text = "Files to Scan")
    progres_label.grid(row = 6, column = 0, pady = 10)

    file_counter_label = tk.Label(window, text = "0")
    file_counter_label.grid(row = 6, column = 1)
    
    scaning_label = tk.Label(window, text = "Scanning files")
    scaning_label.grid(row = 7, column = 1)

    text_area = tk.Text(window, height = 20)
    text_area.grid(row = 8, column = 0, columnspan = 3, padx= 17)

    scrollbar_scanned_books = Scrollbar(window, orient= VERTICAL, relief=SUNKEN,)
    scrollbar_scanned_books.grid(row = 8, column = 4, sticky= N+S)

    text_area.config(yscrollcommand = scrollbar_scanned_books.set)
    scrollbar_scanned_books.config(command = text_area.yview)

    books_error_label = tk.Label(window, text = "Files that could't be read")
    books_error_label.grid(row = 0, column = 5)

    error_text_area = tk.Text(window, width= 50, height= 20, fg= "Red")
    error_text_area.grid(row = 1, rowspan= 4, column = 5, padx = 10)

    scrollbar_error_books = Scrollbar(window, orient= VERTICAL, relief=SUNKEN,)
    scrollbar_error_books.grid(row = 1, rowspan = 4, column = 6, sticky= N+S)

    error_text_area.config(yscrollcommand = scrollbar_error_books.set)
    scrollbar_error_books.config(command = error_text_area.yview)

    pages_error_label = tk.Label(window, text = "Pages that could't be read")
    pages_error_label.grid(row = 6, column = 5)

    pages_error_text_area = tk.Text(window, width= 50, height= 20, fg= "Red")
    pages_error_text_area.grid(row = 6, rowspan= 9, column = 5, padx = 10)

    scrollbar_error_pages = Scrollbar(window, orient= VERTICAL, relief=SUNKEN,)
    scrollbar_error_pages.grid(row = 8, column = 6, sticky= N+S)

    pages_error_text_area.config(yscrollcommand = scrollbar_error_pages.set)
    scrollbar_error_pages.config(command = pages_error_text_area.yview)

    warning_label = tk.Label(window, text = "WARNING: ", fg = "Red")
    warning_label.grid(row = 9, column = 0)

    note_label = tk.Label(window, text = "Sometimes the softaware could say 'Not Responding' \n and it is because it is saving the word document. \n LEAVE IT RUNNING SPECIALLY IF THE DIRECTORY HAS A LOT OF PDFs")
    note_label.grid(row = 9, column = 1)

    window.mainloop()

if __name__ == '__main__':
    interface()
